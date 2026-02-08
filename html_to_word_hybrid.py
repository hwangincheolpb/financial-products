"""
HTML을 Word로 변환: 텍스트는 텍스트로, 이미지는 이미지로
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 문서 생성
doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def add_title(text):
    """제목 추가 (h1)"""
    heading = doc.add_heading(text, level=1)
    heading.runs[0].font.size = Pt(20)
    heading.runs[0].font.bold = True
    heading.paragraph_format.space_after = Pt(12)
    return heading

def add_subtitle(text):
    """부제목 추가 (h2)"""
    heading = doc.add_heading(text, level=2)
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.bold = True
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_text_box(title, items):
    """텍스트 박스 추가"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(3)

def add_image(image_path, width=6.5):
    """이미지 추가"""
    full_path = os.path.join(r"c:\dev\active-projects\fund-comparison-web", image_path)
    if os.path.exists(full_path):
        doc.add_picture(full_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.paragraph_format.space_after = Pt(12)
    else:
        doc.add_paragraph(f"[이미지 없음: {image_path}]")

def add_page_break():
    """페이지 구분"""
    doc.add_page_break()

# ========== PAGE 1: 비교 & 성과 ==========
add_title("타임폴리오 vs 디에스 사모펀드 비교")

add_subtitle("1. 핵심 전략 비교")

add_text_box("타임폴리오 The Time (Long-Short)", [
    "지수 중립 알파 추출 전략. 하락장에서도 수익 추구하는 방어적 절대수익형",
    "10년 누적: +289.3% (KOSPI +131.3%) / 연환산: 15.1% / 샤프지수: 1.26",
    "성과보수: 초과수익의 20% (HWM) / 최소가입: 5억(일반), 1억(전문)"
])

doc.add_paragraph()  # 간격

add_text_box("디에스 Brave S (Long-Biased)", [
    "성장 산업 초입 선점. 지수 대비 압도적 초과수익을 노리는 공격적 수익형",
    "누적: +44.52% (설정 후 4개월) / YTD: +17.93% / 상대 알파: +34.7%p",
    "성과보수: 초과수익의 15% (HWM) / 최소가입: 3억(일반), 1억(전문)"
])

add_subtitle("2. 타임폴리오 10년 성과")
add_image("images/timefolio_performance.png")

add_subtitle("3. 전사 사모펀드 성과 현황")
add_image("images/fund_performance.png")

add_page_break()

# ========== PAGE 2: 타임폴리오 1 ==========
add_title("타임폴리오 The Time - 운용 전략")

add_subtitle("Long/Short 전략 구성")
add_image("images/timefolio_strategy.png")

add_subtitle("운용 철학: 최적의 타이밍")
add_image("images/timefolio_philosophy.png")

add_page_break()

# ========== PAGE 3: 타임폴리오 2 ==========
add_title("타임폴리오 The Time - 운용 시스템")

add_subtitle("MMS (Multi Manager System)")
add_image("images/timefolio_mms.png")

add_subtitle("TMS (Timefolio Management System)")
add_image("images/timefolio_tms.png")

add_page_break()

# ========== PAGE 4: 디에스 1 ==========
add_title("디에스 Brave S - 운용 철학 & 전략")

add_subtitle("운용철학 및 투자스타일")
add_image("images/ds_philosophy.png")

add_subtitle("핵심 운용전략")
add_image("images/ds_core_strategy.png")

add_page_break()

# ========== PAGE 5: 디에스 2 ==========
add_title("디에스 Brave S - 투자 프로세스")

add_subtitle("세부 운용전략")
add_image("images/ds_detail_strategy.png")

add_subtitle("4-Factor 종목선정 프로세스")
add_image("images/ds_4factor.png")

# 문서 저장
output_path = r"c:\dev\active-projects\fund-comparison-web\사모펀드_투자제안서.docx"
doc.save(output_path)
print(f"✅ Word 문서 생성 완료: {output_path}")
